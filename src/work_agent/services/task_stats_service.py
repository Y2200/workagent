"""
任务统计服务（Phase 4）

- get_stats：总览 / 按部门 / 按员工 / 风险任务
- 风险复用 task_reminder_service.compute_risk（确定性，无 LLM）
- to_xlsx / to_docx：导出字节流

多租户：tenant_id=None = 平台全量（SUPER_ADMIN），否则按租户过滤
"""

from datetime import datetime

from io import BytesIO

from work_agent.db.session import SessionLocal
from work_agent.repositories.task_repository import TaskRepository
from work_agent.repositories.user_repository import UserRepository
from work_agent.services.task_reminder_service import task_reminder_service


_OVERDUE_STATUSES = ("pending", "processing")


def _rate(completed: int, total: int) -> float:

    if total <= 0:

        return 0.0

    return round(
        completed * 100.0 / total,
        1,
    )


class TaskStatsService:

    def __init__(
            self,
            repository: TaskRepository | None = None,
            user_repository: UserRepository | None = None
    ):

        self.repository = repository or TaskRepository()

        self.user_repository = user_repository or UserRepository()

    # ======================
    # 统计聚合
    # ======================

    def get_stats(
            self,
            *,
            tenant_id: str | None = None,
            now: datetime | None = None
    ) -> dict:

        now = now or datetime.now()

        db = SessionLocal()

        try:

            tasks = self.repository.list_by_tenant(
                db,
                tenant_id,
            )

        finally:

            db.close()

        # 员工显示名映射
        db = SessionLocal()

        try:

            users = {
                u.id: u
                for u in self.user_repository.list_all(
                    db,
                    tenant_id="",
                )
            }

        finally:

            db.close()

        # 逐任务风险（确定性，复用 Phase 3 规则）
        risk_map = {}

        risky_tasks = []

        for task in tasks:

            risk = task_reminder_service.compute_risk(
                task,
                now=now,
            )

            risk_map[task.id] = risk

            if risk["level"] in ("high", "medium"):

                u = users.get(task.employee_id)

                risky_tasks.append(
                    {
                        "task_id": task.id,
                        "title": task.title,
                        "employee_id": task.employee_id,
                        "username": u.username if u else "",
                        "real_name": u.real_name if u else "",
                        "department": task.department,
                        "progress": task.progress,
                        "priority": task.priority,
                        "status": task.status,
                        "deadline": (
                            task.deadline.isoformat()
                            if task.deadline
                            else None
                        ),
                        "risk_level": risk["level"],
                        "risk_reason": risk["reason"],
                    }
                )

        overview = {
            "total": len(tasks),
            "pending": sum(
                1 for t in tasks
                if t.status == "pending"
            ),
            "processing": sum(
                1 for t in tasks
                if t.status == "processing"
            ),
            "completed": sum(
                1 for t in tasks
                if t.status == "completed"
            ),
            "overdue": sum(
                1 for t in tasks
                if (
                    t.deadline
                    and t.deadline < now
                    and t.status in _OVERDUE_STATUSES
                )
            ),
        }

        # 按部门
        dept_map: dict[str, dict] = {}

        for task in tasks:

            d = (
                task.department
                or "未分配"
            )

            agg = dept_map.setdefault(
                d,
                {
                    "department": d,
                    "total": 0,
                    "completed": 0,
                    "high_risk": 0,
                    "medium_risk": 0,
                },
            )

            agg["total"] += 1

            if task.status == "completed":

                agg["completed"] += 1

            level = risk_map[task.id]["level"]

            if level == "high":

                agg["high_risk"] += 1

            elif level == "medium":

                agg["medium_risk"] += 1

        by_department = []

        for agg in dept_map.values():

            agg["completion_rate"] = _rate(
                agg["completed"],
                agg["total"],
            )

            by_department.append(agg)

        by_department.sort(
            key=lambda r: -r["total"]
        )

        # 按员工
        emp_map: dict[int, dict] = {}

        for task in tasks:

            uid = task.employee_id

            u = users.get(uid)

            agg = emp_map.setdefault(
                uid,
                {
                    "employee_id": uid,
                    "username": u.username if u else "",
                    "real_name": u.real_name if u else "",
                    "department": task.department,
                    "total": 0,
                    "completed": 0,
                    "high_risk": 0,
                    "medium_risk": 0,
                },
            )

            agg["total"] += 1

            if task.status == "completed":

                agg["completed"] += 1

            level = risk_map[task.id]["level"]

            if level == "high":

                agg["high_risk"] += 1

            elif level == "medium":

                agg["medium_risk"] += 1

        by_employee = []

        for agg in emp_map.values():

            agg["completion_rate"] = _rate(
                agg["completed"],
                agg["total"],
            )

            by_employee.append(agg)

        by_employee.sort(
            key=lambda r: -r["total"]
        )

        return {
            "overview": overview,
            "by_department": by_department,
            "by_employee": by_employee,
            "risky_tasks": risky_tasks,
        }

    # ======================
    # 导出
    # ======================

    def to_xlsx(
            self,
            stats: dict
    ) -> bytes:

        from openpyxl import Workbook

        wb = Workbook()

        # 总览
        ws = wb.active

        ws.title = "总览"

        ws.append(
            ["指标", "数值"]
        )

        for key, value in stats["overview"].items():

            ws.append(
                [key, value]
            )

        # 按部门
        ws2 = wb.create_sheet(
            "按部门"
        )

        ws2.append(
            ["部门", "任务数", "已完成", "完成率%", "高风险", "中风险"]
        )

        for row in stats["by_department"]:

            ws2.append(
                [
                    row["department"],
                    row["total"],
                    row["completed"],
                    row["completion_rate"],
                    row["high_risk"],
                    row["medium_risk"],
                ]
            )

        # 按员工
        ws3 = wb.create_sheet(
            "按员工"
        )

        ws3.append(
            ["姓名", "用户名", "部门", "任务数", "已完成", "完成率%", "高风险", "中风险"]
        )

        for row in stats["by_employee"]:

            ws3.append(
                [
                    row["real_name"],
                    row["username"],
                    row["department"],
                    row["total"],
                    row["completed"],
                    row["completion_rate"],
                    row["high_risk"],
                    row["medium_risk"],
                ]
            )

        # 风险任务
        ws4 = wb.create_sheet(
            "风险任务"
        )

        ws4.append(
            ["任务", "负责人", "部门", "进度%", "截止", "风险", "原因"]
        )

        for row in stats["risky_tasks"]:

            ws4.append(
                [
                    row["title"],
                    row["real_name"] or row["username"],
                    row["department"],
                    row["progress"],
                    row["deadline"],
                    row["risk_level"],
                    row["risk_reason"],
                ]
            )

        buf = BytesIO()

        wb.save(buf)

        return buf.getvalue()

    def to_docx(
            self,
            stats: dict
    ) -> bytes:

        from docx import Document

        from docx.shared import Pt

        doc = Document()

        doc.add_heading(
            "任务统计",
            level=1,
        )

        doc.add_heading(
            "一、总览",
            level=2,
        )

        for key, value in stats["overview"].items():

            doc.add_paragraph(
                f"{key}: {value}"
            )

        doc.add_heading(
            "二、按部门",
            level=2,
        )

        table = doc.add_table(
            rows=1,
            cols=6,
        )

        header = table.rows[0].cells

        for i, text in enumerate(
            ["部门", "任务数", "已完成", "完成率%", "高风险", "中风险"]
        ):

            header[i].paragraphs[0].add_run(
                text
            ).bold = True

        for row in stats["by_department"]:

            cells = table.add_row().cells

            values = [
                row["department"],
                row["total"],
                row["completed"],
                row["completion_rate"],
                row["high_risk"],
                row["medium_risk"],
            ]

            for i, value in enumerate(values):

                cells[i].text = str(value)

        doc.add_heading(
            "三、按员工",
            level=2,
        )

        table2 = doc.add_table(
            rows=1,
            cols=8,
        )

        header2 = table2.rows[0].cells

        for i, text in enumerate(
            ["姓名", "用户名", "部门", "任务数", "已完成", "完成率%", "高风险", "中风险"]
        ):

            header2[i].paragraphs[0].add_run(
                text
            ).bold = True

        for row in stats["by_employee"]:

            cells = table2.add_row().cells

            values = [
                row["real_name"],
                row["username"],
                row["department"],
                row["total"],
                row["completed"],
                row["completion_rate"],
                row["high_risk"],
                row["medium_risk"],
            ]

            for i, value in enumerate(values):

                cells[i].text = str(value)

        doc.add_heading(
            "四、风险任务",
            level=2,
        )

        if not stats["risky_tasks"]:

            doc.add_paragraph(
                "当前无高风险/中风险任务"
            )

        for row in stats["risky_tasks"]:

            doc.add_paragraph(
                f"{row['title']}（{row['real_name'] or row['username']}）"
                f"：{row['risk_level']} - {row['risk_reason']}"
            )

        buf = BytesIO()

        doc.save(buf)

        return buf.getvalue()


# 全局单例
task_stats_service = TaskStatsService()
