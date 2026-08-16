"""
任务周报服务（Phase 4）

一份汇总周报：近 7 天完成情况 / 延期任务 / 高风险任务 / 建议
- 风险复用 task_reminder_service.compute_risk（确定性）
- to_docx 生成 Word；generate_weekly 返回 {docx_bytes, summary}
- 预留 department 参数（后续可扩展按部门）

多租户：tenant_id=None = 平台全量（SUPER_ADMIN）
"""

from datetime import datetime, timedelta

from io import BytesIO

from work_agent.db.session import SessionLocal
from work_agent.repositories.task_repository import TaskRepository
from work_agent.repositories.user_repository import UserRepository
from work_agent.services.task_reminder_service import task_reminder_service


_OVERDUE_STATUSES = ("pending", "processing")


class TaskReportService:

    def __init__(
            self,
            repository: TaskRepository | None = None,
            user_repository: UserRepository | None = None
    ):

        self.repository = repository or TaskRepository()

        self.user_repository = user_repository or UserRepository()

    def build_weekly_report(
            self,
            *,
            tenant_id: str | None = None,
            department: str | None = None,
            now: datetime | None = None
    ) -> dict:

        now = now or datetime.now()

        since = now - timedelta(days=7)

        db = SessionLocal()

        try:

            tasks = self.repository.list_by_tenant(
                db,
                tenant_id,
            )

        finally:

            db.close()

        if department:

            tasks = [
                t for t in tasks
                if (t.department or "") == department
            ]

        # 本周完成任务（status=completed 且 updated_at 落在近 7 天）
        completed_this_week = [
            t for t in tasks
            if (
                t.status == "completed"
                and t.updated_at
                and t.updated_at >= since
            )
        ]

        # 延期任务
        overdue_tasks = [
            t for t in tasks
            if (
                t.deadline
                and t.deadline < now
                and t.status in _OVERDUE_STATUSES
            )
        ]

        # 高风险任务（确定性规则）
        risky_tasks = []

        for t in tasks:

            risk = task_reminder_service.compute_risk(
                t,
                now=now,
            )

            if risk["level"] in ("high", "medium"):

                risky_tasks.append(
                    {
                        "task_id": t.id,
                        "title": t.title,
                        "department": t.department,
                        "progress": t.progress,
                        "deadline": (
                            t.deadline.isoformat()
                            if t.deadline
                            else None
                        ),
                        "risk_level": risk["level"],
                        "risk_reason": risk["reason"],
                    }
                )

        high_risk = sum(
            1 for r in risky_tasks
            if r["risk_level"] == "high"
        )

        medium_risk = (
            len(risky_tasks) - high_risk
        )

        # 按部门小计
        dept_count: dict[str, int] = {}

        for t in tasks:

            d = (
                t.department
                or "未分配"
            )

            dept_count[d] = dept_count.get(d, 0) + 1

        by_department = [
            {
                "department": d,
                "total": n,
            }
            for d, n in sorted(
                dept_count.items(),
                key=lambda kv: -kv[1],
            )
        ]

        suggestions = self._build_suggestions(
            high_risk=high_risk,
            overdue=len(overdue_tasks),
            completed=len(completed_this_week),
            total=len(tasks),
            risky_tasks=risky_tasks,
        )

        return {
            "generated_at": now.isoformat(),
            "period": {
                "start": since.isoformat(),
                "end": now.isoformat(),
            },
            "summary": {
                "total": len(tasks),
                "completed_this_week": len(completed_this_week),
                "overdue": len(overdue_tasks),
                "high_risk": high_risk,
                "medium_risk": medium_risk,
            },
            "completed_tasks": [
                t.title
                for t in completed_this_week
            ],
            "overdue_tasks": [
                {
                    "title": t.title,
                    "deadline": (
                        t.deadline.isoformat()
                        if t.deadline
                        else None
                    ),
                    "progress": t.progress,
                }
                for t in overdue_tasks
            ],
            "risky_tasks": risky_tasks,
            "by_department": by_department,
            "suggestions": suggestions,
        }

    @staticmethod
    def _build_suggestions(
            *,
            high_risk: int,
            overdue: int,
            completed: int,
            total: int,
            risky_tasks: list[dict]
    ) -> list[str]:

        suggestions = []

        if high_risk > 0:

            dept = "、".join(
                sorted(
                    {
                        r["department"]
                        for r in risky_tasks
                        if r["risk_level"] == "high"
                    }
                )
            )

            suggestions.append(
                f"{high_risk} 个高风险任务（涉及部门：{dept or '未分配'}），"
                "建议优先跟进处理。"
            )

        if overdue > 0:

            suggestions.append(
                f"{overdue} 个任务已逾期，建议督促负责人尽快补齐进度。"
            )

        if completed == 0 and total > 0:

            suggestions.append(
                "本周暂无完成任务，建议关注进行中任务的推进情况。"
            )

        if not suggestions:

            suggestions.append(
                "本周任务整体正常，继续保持。"
            )

        return suggestions

    def to_docx(
            self,
            report: dict
    ) -> bytes:

        from docx import Document

        doc = Document()

        doc.add_heading(
            "任务周报",
            level=1,
        )

        doc.add_paragraph(
            f"统计周期：{report['period']['start'][:10]} ~ "
            f"{report['period']['end'][:10]}"
        )

        doc.add_heading(
            "一、本周完成情况",
            level=2,
        )

        s = report["summary"]

        doc.add_paragraph(
            f"本周完成任务：{s['completed_this_week']} 个"
        )

        doc.add_paragraph(
            f"当前进行中任务：{s['total'] - s['completed_this_week']} 个"
        )

        doc.add_heading(
            "二、延期任务",
            level=2,
        )

        if not report["overdue_tasks"]:

            doc.add_paragraph(
                "本周无延期任务"
            )

        for task in report["overdue_tasks"]:

            doc.add_paragraph(
                f"{task['title']}（截止 {task['deadline']}，进度 {task['progress']}%）"
            )

        doc.add_heading(
            "三、高风险任务",
            level=2,
        )

        if not report["risky_tasks"]:

            doc.add_paragraph(
                "本周无高风险任务"
            )

        for task in report["risky_tasks"]:

            doc.add_paragraph(
                f"{task['title']}（{task['risk_level']}）：{task['risk_reason']}"
            )

        doc.add_heading(
            "四、建议",
            level=2,
        )

        for tip in report["suggestions"]:

            doc.add_paragraph(
                f"- {tip}"
            )

        buf = BytesIO()

        doc.save(buf)

        return buf.getvalue()

    def generate_weekly(
            self,
            *,
            tenant_id: str | None = None,
            department: str | None = None,
            now: datetime | None = None
    ) -> dict:

        """
        生成周报：返回 {docx_bytes, summary}
        """

        report = self.build_weekly_report(
            tenant_id=tenant_id,
            department=department,
            now=now,
        )

        return {
            "docx_bytes": self.to_docx(report),
            "summary": report,
        }

    # ======================
    # 周报部门经理投递（Enterprise Agent Phase 4）
    # ======================

    def send_department_digests(
            self,
            *,
            tenant_id: str | None = None,
            now: datetime | None = None
    ) -> dict:

        """
        按部门聚合周报摘要 → 企微 digest 发给该部门经理（DEPARTMENT_ADMIN）

        部门经理看到的是其所在部门所有员工的汇总，不是逐条推送。
        dry 模式下只统计。

        返回：
        {departments, digests_sent, skipped_unbound, failed}
        """

        from work_agent.db.models import User
        from work_agent.repositories.rbac_repository import RBACRepository
        from work_agent.services.notification_service import notification_service

        now = now or datetime.now()

        db = SessionLocal()

        try:

            # 1. 全部部门管理员
            admin_ids = RBACRepository().list_user_ids_by_role(
                db,
                "DEPARTMENT_ADMIN",
            )

            if not admin_ids:

                return {
                    "departments": [],
                    "digests_sent": 0,
                    "skipped_unbound": 0,
                    "failed": 0,
                }

            admins = (
                db.query(User)
                .filter(User.id.in_(admin_ids))
                .all()
            )

            # 2. 每个部门管理员 → 本部门周报摘要
            summary = {
                "departments": [],
                "digests_sent": 0,
                "skipped_unbound": 0,
                "failed": 0,
            }

            seen_departments = set()

            for admin in admins:

                dept = admin.department or ""

                if not dept:
                    continue

                if dept in seen_departments:
                    continue

                seen_departments.add(dept)

                report = self.build_weekly_report(
                    tenant_id=tenant_id,
                    department=dept,
                    now=now,
                )

                content = self._department_digest_text(
                    dept,
                    report,
                )

                summary["departments"].append(dept)

                # 发给该部门所有部门管理员（聚合，非逐条）
                for mgr in admins:

                    if (mgr.department or "") != dept:
                        continue

                    if not mgr.wechat_user_id:

                        notification_service.record(
                            tenant_id=(
                                mgr.tenant_id
                                or tenant_id
                                or ""
                            ),
                            task_id=0,
                            receiver_id=mgr.id,
                            channel="wechat",
                            content=content,
                            status="failed",
                        )

                        summary["skipped_unbound"] += 1

                        continue

                    result = notification_service.send_wechat(
                        tenant_id=(
                            mgr.tenant_id
                            or tenant_id
                            or ""
                        ),
                        task_id=0,
                        receiver_id=mgr.id,
                        wechat_user_id=mgr.wechat_user_id,
                        content=content,
                    )

                    if result.get("ok"):

                        summary["digests_sent"] += 1

                    else:

                        summary["failed"] += 1

            return summary

        finally:

            db.close()

    @staticmethod
    def _department_digest_text(
            department: str,
            report: dict
    ) -> str:

        """
        部门周报摘要文案（发给部门经理）
        """

        completed = report.get("completed_tasks") or []

        overdue = report.get("overdue_tasks") or []

        risky = report.get("risky_tasks") or []

        lines = [
            f"【周报】{department} 本周工作汇总",
            "",
            f"完成任务：{len(completed)} 个",
            f"延期任务：{len(overdue)} 个",
            f"高风险任务：{len(risky)} 个",
        ]

        if completed:

            lines.append("")
            lines.append("完成：")

            for title in completed[:5]:

                lines.append(f"- {title}")

        if risky:

            lines.append("")
            lines.append("重点关注：")

            for t in risky[:5]:

                lines.append(
                    f"- {t['title']}"
                    f"（进度{t['progress']}%，{t['risk_reason']}）"
                )

        lines.append("")
        lines.append("详情请登录 Web 后台查看周报。")

        return "\n".join(lines)


# 全局单例
task_report_service = TaskReportService()
