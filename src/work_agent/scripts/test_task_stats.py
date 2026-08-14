"""
任务统计 / 周报 / 邮件测试（Phase 4）

Part 1  统计聚合（overview / 部门 / 员工 / 完成率 / 风险计数）
Part 2  多租户隔离（租户管理员仅本租户）
Part 3  导出 xlsx / docx 非空 + 读回校验
Part 4  周报 build + to_docx + 建议
Part 5  EmailService（mock smtplib：成功 / 失败 / 未配置 skipped）
Part 6  任务完成邮件钩子（confirm 到 100% → mock 发邮件给创建者 + 记 channel=email）

用法：
    python -m work_agent.scripts.test_task_stats
"""

from datetime import datetime, timedelta

from io import BytesIO

from work_agent.config import settings
from work_agent.db.models import User
from work_agent.db.models.task import Task, TaskNotification
from work_agent.db.session import SessionLocal
from work_agent.repositories.task_repository import TaskRepository
from work_agent.repositories.user_repository import UserRepository
from work_agent.scripts.migrate_user_profile import migrate as migrate_user_profile
from work_agent.scripts.seed_rbac import seed_rbac
from work_agent.scripts.seed_tenants import seed_tenants
from work_agent.services.auth_service import AuthService
from work_agent.services.task_report_service import task_report_service
from work_agent.services.task_stats_service import task_stats_service
from work_agent.services.task_service import task_service


_TEST_USERS = [
    "统计创建者",
    "统计员工",
]

_CREATED_TASKS: list[int] = []


def _setup():

    seed_tenants()

    seed_rbac()

    migrate_user_profile()

    _cleanup()


def _cleanup():

    db = SessionLocal()

    try:

        # 清空测试租户（1/2）的任务与通知，保证统计确定性（测试脚本数据可丢弃）
        db.query(TaskNotification).filter(
            TaskNotification.tenant_id.in_(["1", "2"])
        ).delete(
            synchronize_session=False
        )

        db.query(Task).filter(
            Task.tenant_id.in_(["1", "2"])
        ).delete(
            synchronize_session=False
        )

        for name in _TEST_USERS:

            user = UserRepository().get_by_username(
                db,
                name,
            )

            if user:

                db.delete(user)

        db.commit()

    finally:

        db.close()


def _user(username):

    db = SessionLocal()

    try:

        return UserRepository().get_by_username(
            db,
            username,
        )

    finally:

        db.close()


def _create_task(
        employee,
        *,
        title: str,
        department: str = "",
        deadline=None,
        progress: int = 0,
        priority: str = "normal",
        creator_id=None
):

    db = SessionLocal()

    try:

        task = task_service.create_task(
            creator_tenant_id=employee.tenant_id,
            title=title,
            creator_id=creator_id or employee.id,
            employee_id=employee.id,
            department=department,
            deadline=deadline,
            priority=priority,
        )

        if progress:

            TaskRepository().update_progress(
                db,
                task.id,
                progress,
            )

        _CREATED_TASKS.append(task.id)

        return task

    finally:

        db.close()


# ======================
# Part 1 统计聚合
# ======================

def test_stats_aggregation():

    emp = _user("A财务员工")

    assert emp, "需要 A财务员工 测试用户"

    now = datetime.now()

    _create_task(
        emp,
        title="逾期研发",
        department="研发部",
        deadline=now - timedelta(days=2),
        progress=30,
    )

    _create_task(
        emp,
        title="正常研发",
        department="研发部",
        deadline=now + timedelta(days=10),
        progress=80,
    )

    _create_task(
        emp,
        title="完成财务",
        department="财务部",
        deadline=now + timedelta(days=5),
        progress=100,
    )

    _create_task(
        emp,
        title="临期财务",
        department="财务部",
        deadline=now + timedelta(days=3),
        progress=75,
    )

    stats = task_stats_service.get_stats(
        tenant_id=emp.tenant_id,
        now=now,
    )

    ov = stats["overview"]

    assert ov["total"] == 4, ov

    assert ov["completed"] == 1, ov

    assert ov["processing"] == 3, ov

    assert ov["overdue"] == 1, ov

    depts = {
        row["department"]: row
        for row in stats["by_department"]
    }

    assert depts["研发部"]["total"] == 2, depts

    assert depts["研发部"]["completed"] == 0, depts

    assert depts["研发部"]["completion_rate"] == 0.0, depts

    assert depts["研发部"]["high_risk"] == 1, depts

    assert depts["财务部"]["total"] == 2, depts

    assert depts["财务部"]["completed"] == 1, depts

    assert depts["财务部"]["completion_rate"] == 50.0, depts

    assert depts["财务部"]["medium_risk"] == 1, depts

    assert len(stats["by_employee"]) == 1, stats["by_employee"]

    emp_row = stats["by_employee"][0]

    assert emp_row["total"] == 4, emp_row

    assert emp_row["completion_rate"] == 25.0, emp_row

    risky = stats["risky_tasks"]

    assert len(risky) == 2, risky

    levels = {
        r["risk_level"]
        for r in risky
    }

    assert {"high", "medium"} <= levels, risky

    print("Part 1 ✅ 统计聚合（overview/部门/员工/完成率/风险计数）")


# ======================
# Part 2 多租户隔离
# ======================

def test_tenant_isolation():

    emp_a = _user("A财务员工")

    emp_b = _user("B市场员工")

    assert emp_b, "需要 B市场员工 测试用户"

    now = datetime.now()

    _create_task(
        emp_b,
        title="B租户任务",
        department="市场部",
        deadline=now - timedelta(days=1),
        progress=20,
    )

    stats_a = task_stats_service.get_stats(
        tenant_id=emp_a.tenant_id,
        now=now,
    )

    depts_a = {
        row["department"]
        for row in stats_a["by_department"]
    }

    assert "市场部" not in depts_a, depts_a

    stats_b = task_stats_service.get_stats(
        tenant_id=emp_b.tenant_id,
        now=now,
    )

    assert stats_b["overview"]["total"] == 1, stats_b

    print("Part 2 ✅ 多租户隔离（租户管理员仅本租户）")


# ======================
# Part 3 导出
# ======================

def test_export():

    emp = _user("A财务员工")

    stats = task_stats_service.get_stats(
        tenant_id=emp.tenant_id,
    )

    xlsx = task_stats_service.to_xlsx(stats)

    assert xlsx and len(xlsx) > 0, "xlsx 应非空"

    from openpyxl import load_workbook

    wb = load_workbook(BytesIO(xlsx))

    assert "总览" in wb.sheetnames, wb.sheetnames

    docx = task_stats_service.to_docx(stats)

    assert docx and len(docx) > 0, "docx 应非空"

    from docx import Document as DocxDocument

    doc = DocxDocument(BytesIO(docx))

    text = "\n".join(
        p.text
        for p in doc.paragraphs
    )

    assert "任务统计" in text, text

    print("Part 3 ✅ 导出 xlsx/docx（非空 + 读回校验）")


# ======================
# Part 4 周报
# ======================

def test_weekly_report():

    emp = _user("A财务员工")

    now = datetime.now()

    report = task_report_service.build_weekly_report(
        tenant_id=emp.tenant_id,
        now=now,
    )

    assert report["summary"]["completed_this_week"] == 1, report["summary"]

    assert report["summary"]["overdue"] == 1, report["summary"]

    assert report["summary"]["high_risk"] == 1, report["summary"]

    assert report["suggestions"], report["suggestions"]

    docx = task_report_service.to_docx(report)

    assert docx and len(docx) > 0, "周报 docx 应非空"

    from docx import Document as DocxDocument

    doc = DocxDocument(BytesIO(docx))

    text = "\n".join(
        p.text
        for p in doc.paragraphs
    )

    assert "任务周报" in text, text

    assert "建议" in text, text

    print("Part 4 ✅ 周报（build + to_docx + 建议）")


# ======================
# Part 5 EmailService
# ======================

def test_email_service():

    import work_agent.services.email_service as es

    saved = (
        settings.email_enabled,
        settings.smtp_host,
        settings.smtp_port,
        settings.smtp_username,
    )

    try:

        # ① 未启用 → skipped
        settings.email_enabled = False

        settings.smtp_host = "smtp.test"

        r = es.email_service.send(
            to="a@x.com",
            subject="s",
            content="c",
        )

        assert r["status"] == "skipped", r

        # ② 启用但未配置 host → skipped
        settings.email_enabled = True

        settings.smtp_host = ""

        r = es.email_service.send(
            to="a@x.com",
            subject="s",
            content="c",
        )

        assert r["status"] == "skipped", r

        # ③ 配置完整 → 发送成功
        settings.smtp_host = "smtp.test"

        settings.smtp_port = 465

        settings.smtp_username = "user@test.com"

        real_smtp_ssl = es.smtplib.SMTP_SSL

        class FakeSMTP:

            def __init__(self, *a, **kw):

                self.sent = []

            def login(self, u, p):

                pass

            def sendmail(self, f, t, m):

                self.sent.append((f, t))

            def quit(self):

                pass

        es.smtplib.SMTP_SSL = FakeSMTP

        r = es.email_service.send(
            to="a@x.com",
            subject="s",
            content="c",
        )

        assert r["status"] == "sent", r

        # ④ SMTP 抛异常 → failed（不抛给调用方）
        class FailSMTP:

            def __init__(self, *a, **kw):

                raise Exception("smtp down")

        es.smtplib.SMTP_SSL = FailSMTP

        r = es.email_service.send(
            to="a@x.com",
            subject="s",
            content="c",
        )

        assert r["status"] == "failed", r

        es.smtplib.SMTP_SSL = real_smtp_ssl

    finally:

        (
            settings.email_enabled,
            settings.smtp_host,
            settings.smtp_port,
            settings.smtp_username,
        ) = saved

    print("Part 5 ✅ EmailService（未配置 skipped / 成功 sent / 异常 failed）")


# ======================
# Part 6 任务完成邮件钩子
# ======================

def test_task_completed_email():

    import work_agent.services.email_service as es

    saved = (
        settings.email_enabled,
        settings.smtp_host,
        settings.smtp_port,
        settings.smtp_username,
    )

    try:

        settings.email_enabled = True

        settings.smtp_host = "smtp.test"

        settings.smtp_port = 465

        settings.smtp_username = "user@test.com"

        real_smtp_ssl = es.smtplib.SMTP_SSL

        class FakeSMTP:

            instances = []

            def __init__(self, *a, **kw):

                FakeSMTP.instances.append(self)

                self.sent = []

            def login(self, u, p):

                pass

            def sendmail(self, f, t, m):

                self.sent.append((f, t))

            def quit(self):

                pass

        es.smtplib.SMTP_SSL = FakeSMTP

        # 创建者（带邮箱）+ 员工
        db = SessionLocal()

        try:

            creator = UserRepository().create(
                db,
                username="统计创建者",
                password_hash=AuthService.hash_password("test123"),
                department="研发部",
                role="管理员",
                email="creator@test.com",
                tenant_id="1",
            )

            emp = UserRepository().create(
                db,
                username="统计员工",
                password_hash=AuthService.hash_password("test123"),
                department="研发部",
                role="员工",
                real_name="统计员工",
                tenant_id="1",
            )

        finally:

            db.close()

        now = datetime.now()

        task = _create_task(
            emp,
            title="邮件完成测试",
            department="研发部",
            deadline=now + timedelta(days=5),
            creator_id=creator.id,
        )

        # 提交 → 确认 100% → 完成 → 发邮件
        db = SessionLocal()

        try:

            TaskRepository().upsert_pending(
                db,
                task_id=task.id,
                employee_id=emp.id,
                content="全部完成",
                parsed={
                    "progress": 100,
                    "summary": "全部完成",
                },
            )

        finally:

            db.close()

        conf = task_service.confirm_pending(
            tenant_id=emp.tenant_id,
            employee_id=emp.id,
        )

        assert conf["status"] == "confirmed", conf

        sent_calls = [
            call
            for inst in FakeSMTP.instances
            for call in inst.sent
        ]

        assert sent_calls, "应向创建者发送邮件"

        assert any(
            "creator@test.com" in call[1]
            for call in sent_calls
        ), sent_calls

        # 落库 task_notifications channel=email
        rows = _notifications(task.id)

        email_rows = [
            r
            for r in rows
            if r.channel == "email"
        ]

        assert email_rows, rows

        assert email_rows[0].status == "sent", email_rows[0].status

        es.smtplib.SMTP_SSL = real_smtp_ssl

    finally:

        (
            settings.email_enabled,
            settings.smtp_host,
            settings.smtp_port,
            settings.smtp_username,
        ) = saved

    print("Part 6 ✅ 任务完成邮件钩子（confirm 100% → 发邮件 + 记 channel=email）")


def _notifications(task_id):

    db = SessionLocal()

    try:

        return (
            db.query(TaskNotification)
            .filter(TaskNotification.task_id == task_id)
            .all()
        )

    finally:

        db.close()


def test():

    _setup()

    test_stats_aggregation()

    test_tenant_isolation()

    test_export()

    test_weekly_report()

    test_email_service()

    test_task_completed_email()


if __name__ == "__main__":

    test()
