import io

from dataclasses import dataclass

from pypdf import PdfReader
from docx import Document as DocxDocument

from work_agent.rag.loader import parse_markdown_metadata


SUPPORTED_EXTENSIONS = {
    ".pdf",
    ".docx",
    ".md",
    ".txt"
}


@dataclass
class ParsedDocument:

    """
    解析结果
    """

    filename: str

    content: str

    metadata: dict

    file_type: str


def parse_document(
        filename: str,
        data: bytes
) -> ParsedDocument:

    """
    解析上传文档为纯文本

    支持 pdf / docx / md / txt
    """

    file_type = (
        filename.rsplit(".", 1)[-1].lower()
        if "." in filename
        else ""
    )

    extension = f".{file_type}"

    if extension not in SUPPORTED_EXTENSIONS:

        raise ValueError(
            f"不支持的文件类型: {extension}"
        )


    if file_type == "pdf":

        content = _parse_pdf(data)

        metadata = {}

    elif file_type == "docx":

        content = _parse_docx(data)

        metadata = {}

    elif file_type == "md":

        text = _decode(data)

        metadata, content = parse_markdown_metadata(
            text
        )

    else:

        content = _decode(data)

        metadata = {}


    if not content.strip():

        raise ValueError(
            "文档无文本内容"
        )


    return ParsedDocument(
        filename=filename,
        content=content,
        metadata=metadata,
        file_type=file_type
    )


def _decode(data: bytes) -> str:

    """
    utf-8 优先，失败回退 gbk
    """

    try:
        return data.decode("utf-8")

    except UnicodeDecodeError:
        return data.decode(
            "gbk",
            errors="ignore"
        )


def _parse_pdf(data: bytes) -> str:

    """
    文本型 PDF 逐页提取
    """

    reader = PdfReader(
        io.BytesIO(data)
    )

    pages = []

    for page in reader.pages:

        text = page.extract_text() or ""

        pages.append(text)

    return "\n".join(pages)


def _parse_docx(data: bytes) -> str:

    """
    python-docx 段落拼接
    """

    document = DocxDocument(
        io.BytesIO(data)
    )

    paragraphs = [
        paragraph.text
        for paragraph in document.paragraphs
        if paragraph.text
    ]

    return "\n".join(paragraphs)
